from flask import Flask, request, jsonify, send_from_directory, send_file, render_template
import pdfplumber
import pytesseract
from PIL import Image
import os
import openai
from docx import Document
import fitz  # PyMuPDF
import io

# Desenvolvido por Levi Lucena - linkedin.com/in/levilucena

app = Flask(__name__)

# Configuração de diretórios
UPLOAD_FOLDER = 'static/images'
CERTIFICATE_FOLDER = 'certificates'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CERTIFICATE_FOLDER, exist_ok=True)

# Configuração do caminho do Tesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Configuração da chave da API OpenAI
openai.api_key = 'SUA CHAVE API AQUI'  # INSIRA A CHAVE API

def extract_text_from_pdf(file):
    text = ""
    images = []

    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text
            else:
                # OCR como fallback para imagens
                image = page.to_image().original
                text += ocr_from_image(image)

    # Voltar ponteiro do arquivo para PyMuPDF
    file.seek(0)

    try:
        pdf_document = fitz.open(stream=file.read(), filetype='pdf')
        if len(pdf_document) > 0:
            page = pdf_document.load_page(0)
            for img_index, img in enumerate(page.get_images(full=True)):
                base_image = pdf_document.extract_image(img[0])
                image_bytes = base_image["image"]
                image = Image.open(io.BytesIO(image_bytes))
                image_filename = f'image_0_{img_index}.png'
                image_path = os.path.join(UPLOAD_FOLDER, image_filename)
                image.save(image_path)
                images.append(image_filename)
    except Exception as e:
        print(f"Erro ao extrair imagens: {e}")

    return text, images

def ocr_from_image(image):
    return pytesseract.image_to_string(image)

def extract_text_from_docx(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def extract_text_from_txt(file):
    return file.read().decode('utf-8')

def generate_certificate(fields):
    template_path = 'path_to_certificate_template.docx'
    doc = Document(template_path)

    placeholders = {
        "EMPRESA CONTRATANTE": fields.get('company_name', 'N/A'),
        "ENDEREÇO DO CONTRATANTE": fields.get('contract_address', 'N/A'),
        "CNPJ EMPRESA CONTRATANTE": fields.get('cnpj_company', 'N/A'),
        "DATA DE ASSINATURA DO CONTRATO": fields.get('contract_date', 'N/A'),
        "NÚMERO DO PREGÃO ELETRÔNICO": fields.get('auction_number', 'N/A'),
        "NÚMERO DO CONTRATO": fields.get('contract_number', 'N/A'),
        "OBJETIVO SOLICITADO NO CONTRATO": fields.get('objective', 'N/A'),
        "DATA DE INÍCIO DO CONTRATO": fields.get('contract_start', 'N/A'),
        "DATA DE TÉRMINO DO CONTRATO": fields.get('contract_end', 'N/A'),
        "SERVIÇOS PRESTADOS": fields.get('services_provided', 'N/A'),
        "LOCAL DE EXECUÇÃO DOS SERVIÇOS": fields.get('location_of_services', 'N/A'),
        "INFORMAR VOLUMETRIA DOS SERVIÇOS PRESTADOS": fields.get('volumetry', 'N/A'),
        "EQUIPE OPERACIONAL": fields.get('operational_team', 'N/A'),
        "AMBIENTE COMPUTACIONAL": fields.get('computing_environment', 'N/A'),
        "FERRAMENTAS UTILIZADAS": fields.get('tools_used', 'N/A'),
        "TECNOLOGIAS ENVOLVIDAS": fields.get('technologies_involved', 'N/A'),
        "METOLOGIAS": fields.get('methodology', 'N/A'),
        "NÍVEIS DE SERVIÇO (SLA)": fields.get('methodology', 'N/A'),
        "SATISFAÇÃO": fields.get('satisfaction', 'N/A'),
        "Nome": fields.get('signatory_name', 'N/A'),
        "Cargo": fields.get('signatory_role', 'N/A'),
        "E-Mail": fields.get('signatory_email', 'N/A'),
        "Telefone": fields.get('signatory_phone', 'N/A'),
        "LOCAL": fields.get('location', 'N/A'),
        "DATA": fields.get('date', 'N/A')
    }

    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    output_path = os.path.join(CERTIFICATE_FOLDER, 'certificado.docx')
    doc.save(output_path)

    return output_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return jsonify({'error': 'Nenhum arquivo enviado'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nenhum arquivo selecionado'}), 400

    if file and file.filename.endswith('.pdf'):
        text, images = extract_text_from_pdf(file)
        return jsonify({'text': text, 'images': images})
    
    elif file.filename.endswith('.docx'):
        text = extract_text_from_docx(file)
        return jsonify({'text': text})

    elif file.filename.endswith('.txt'):
        text = extract_text_from_txt(file)
        return jsonify({'text': text})

    return jsonify({'error': 'Tipo de arquivo inválido'}), 400

@app.route('/ask', methods=['POST'])
def ask():
    data = request.json
    question = data.get('question')
    document_text = data.get('document_text')

    if not document_text:
        return jsonify({'error': 'Texto do documento ausente'}), 400

    try:
        # Prompt fixo
        prompt = f"""
        Você tem a tarefa de preencher um documento com base nos dados extraídos. O documento contém placeholders que precisam ser substituídos pelos valores correspondentes extraídos.

        Dados Extraídos:
        {document_text}

        Tarefa:
            Identificar e extrair os seguintes campos do texto e formatar a saída da seguinte forma:

            EMPRESA CONTRATANTE: [Placeholder]
            ENDEREÇO DO CONTRATANTE: [Placeholder]
            CNPJ EMPRESA CONTRATANTE: [Placeholder]
            DATA DE ASSINATURA DO CONTRATO: [Placeholder]
            NÚMERO DO PREGÃO ELETRÔNICO: [Placeholder]
            NÚMERO DO CONTRATO: [Placeholder]
            OBJETIVO SOLICITADO NO CONTRATO: [Placeholder]
            DATA DE INÍCIO DO CONTRATO: [Placeholder]
            DATA DE TÉRMINO DO CONTRATO: [Placeholder]
            SERVIÇOS PRESTADOS: [Placeholder]
            LOCAL DE EXECUÇÃO DOS SERVIÇOS: [Placeholder]
            INFORMAR VOLUMETRIA DOS SERVIÇOS PRESTADOS: [Placeholder]
            EQUIPE OPERACIONAL: [Placeholder]
            AMBIENTE COMPUTACIONAL: [Placeholder]
            FERRAMENTAS UTILIZADAS: [Placeholder]
            TECNOLOGIAS ENVOLVIDAS: [Placeholder]
            METODOLOGIAS: [Placeholder]
            NÍVEIS DE SERVIÇO (SLA): [Placeholder]
            SATISFAÇÃO: [Placeholder]
            Nome: [Placeholder]
            Cargo: [Placeholder]
            E-Mail: [Placeholder]
            Telefone: [Placeholder]
            LOCAL: [Placeholder]
            DATA: [Placeholder]
            ANO: [Placeholder]

            Por favor, retorne a resposta no formato acima, substituindo [Placeholder] pelos valores extraídos do texto. Cada campo deve estar em uma linha separada.
        """

        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",  # Ou "gpt-4" ou "gpt-3.5-turbo" dependendo da sua escolha
            messages=[
                {"role": "system", "content": "Você é um assistente útil."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.7
        )

        extracted_info = response.choices[0].message['content'].strip()
        formatted_info = '\n'.join(line.strip() for line in extracted_info.split('\n') if line.strip())

        return jsonify({'extracted_info': formatted_info})
    except Exception as e:
        return jsonify({'error': f'Error from OpenAI: {str(e)}'}), 500

@app.route('/generate_certificate', methods=['POST'])
def generate_certificate_route():
    data = request.json
    fields = {
        'contract_address': data.get('contract_address'),
        'contract_date': data.get('contract_date'),
        'auction_number': data.get('auction_number'),
        'contract_start': data.get('contract_start'),
        'contract_end': data.get('contract_end'),
        'volumetry': data.get('volumetry'),
        'signatory_name': data.get('signatory_name'),
        'signatory_role': data.get('signatory_role'),
        'signatory_email': data.get('signatory_email'),
        'signatory_phone': data.get('signatory_phone'),
        'location': data.get('location'),
        'date': data.get('date')
    }

    try:
        certificate_path = generate_certificate(fields)
        return jsonify({'message': 'Certificate generated successfully', 'certificate_path': certificate_path})
    except Exception as e:
        return jsonify({'error': f'Error generating certificate: {e}'}), 500

@app.route('/generate_docx', methods=['POST'])
def generate_docx():
    data = request.json
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    # Crie o documento
    doc = Document()

    # Adicione o conteúdo formatado
    doc.add_heading('ATESTADO DE CAPACIDADE TÉCNICA', level=1)

    # Adicione os parágrafos com os dados
    doc.add_paragraph(
        f"{data.get('EMPRESA CONTRATANTE', '__________')}, localizada no {data.get('ENDEREÇO DO CONTRATANTE', '__________')}, "
        f"inscrita no CNPJ sob nº {data.get('CNPJ EMPRESA CONTRATANTE', '__________')}, declara que mantém com a empresa STEFANINI CONSULTORIA E ASSESSORIA EM INFORMÁTICA S.A., "
        f"com sede em Jaguariúna/SP, na Rua Minas Gerais, nº 1.476, Jardim Alice, inscrita no CNPJ nº. 58.069.360/0001-20, o Contrato {data.get('NÚMERO DO CONTRATO', '__________')}, "
        f"firmado em {data.get('DATA DE ASSINATURA DO CONTRATO', '__________')} conforme abaixo discriminado:"
    )
    doc.add_paragraph(
        f"Pregão Eletrônico: {data.get('NÚMERO DO PREGÃO ELETRÔNICO', '__________')}. "
        f"Número do Contrato: {data.get('NÚMERO DO CONTRATO', '__________')}."
    )
    doc.add_paragraph(
        f"Objeto: {data.get('OBJETIVO SOLICITADO NO CONTRATO', '__________')}"
    )
    doc.add_paragraph(
        f"Vigência: {data.get('DATA DE INÍCIO DO CONTRATO', '__________')} a {data.get('DATA DE TÉRMINO DO CONTRATO', '__________')}."
    )
    doc.add_paragraph(
        f"DETALHAMENTO DOS SERVIÇOS PRESTADOS:\n"
        f"O escopo dos serviços prestados encontra-se detalhado no Contrato {data.get('NÚMERO DO CONTRATO', '__________')}, no Edital do Pregão Eletrônico Nº {data.get('NÚMERO DO PREGÃO ELETRÔNICO', '__________')}, "
        f"e em seus anexos e aditivos.\n\n"
        f"{data.get('SERVIÇOS PRESTADOS', '__________')}."
    )
    doc.add_paragraph(
        f"LOCAL DE EXECUÇÃO DOS SERVIÇOS:\n"
        f"{data.get('LOCAL DE EXECUÇÃO DOS SERVIÇOS', '__________')}"
    )
    doc.add_paragraph(
        f"VOLUMETRIA:\n"
        f"{data.get('INFORMAR VOLUMETRIA DOS SERVIÇOS PRESTADOS', '__________')}"
    )
    doc.add_paragraph(
        f"EQUIPE OPERACIONAL:\n"
        f"{data.get('EQUIPE OPERACIONAL', '__________')}"
    )
    doc.add_paragraph(
        f"AMBIENTE COMPUTACIONAL:\n"
        f"{data.get('AMBIENTE COMPUTACIONAL', '__________')}"
    )
    doc.add_paragraph(
        f"FERRAMENTAS UTILIZADAS:\n"
        f"{data.get('FERRAMENTAS UTILIZADAS', '__________')}"
    )
    doc.add_paragraph(
        f"TECNOLOGIAS ENVOLVIDAS:\n"
        f"{data.get('TECNOLOGIAS ENVOLVIDAS', '__________')}"
    )
    doc.add_paragraph(
        f"METODOLOGIAS:\n"
        f"{data.get('METODOLOGIAS', '__________')}"
    )
    doc.add_paragraph(
        f"NÍVEIS DE SERVIÇO (SLA):\n"
        f"{data.get('NÍVEIS DE SERVIÇO (SLA)', '__________')}"
    )
    doc.add_paragraph(
        f"SATISFAÇÃO\n"
        f"{data.get('SATISFAÇÃO', '__________')}"
    )
    doc.add_paragraph(
        f"SIGNATÁRIO\n"
        f"Nome: {data.get('Nome', 'PREENCHER')}\n"
        f"Cargo: {data.get('Cargo', 'PREENCHER')}\n"
        f"E-Mail: {data.get('E-Mail', 'PREENCHER')}\n"
        f"Telefone: {data.get('Telefone', 'PREENCHER')}\n"
        f"{data.get('LOCAL', 'PREENCHER')}, {data.get('DATA', 'PREENCHER')}"
    )

    # Salve o documento em um objeto de bytes
    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    return send_file(
        byte_io,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='certificado.docx'
    )

@app.route('/ask_custom', methods=['POST'])
def ask_custom():
    data = request.json
    question = data.get('question')
    document_text = data.get('document_text')

    if not question or not document_text:
        return jsonify({'error': 'Pergunta ou texto do documento ausente'}), 400

    prompt = f"Documento: {document_text}\n\nPergunta: {question}\n\nResponda com base nas informações do documento."

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",  # Ou "gpt-4o-mini", "gpt-4" ou "gpt-3.5-turbo" dependendo da sua escolha
            messages=[
                {"role": "system", "content": "Você é um assistente útil."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.7
        )
        return jsonify({'answer': response.choices[0].text})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/certificates/<filename>')
def serve_certificate(filename):
    return send_from_directory(CERTIFICATE_FOLDER, filename)

@app.route('/static/images/<filename>')
def serve_image(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)

if __name__ == '__main__':
    app.run(debug=True)

if __name__ == '__main__':
    app.run(debug=True)
